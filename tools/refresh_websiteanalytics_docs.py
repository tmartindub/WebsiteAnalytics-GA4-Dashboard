from __future__ import annotations

import json
import math
import os
import re
import shutil
from datetime import date
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
GUIDES_DIR = ROOT / "docs" / "guides"
PDF_DIR = ROOT / "docs" / "pdf"
HELP_DIR = ROOT / "help"
ASSET_DIR = GUIDES_DIR / "doc-assets"
HELP_ASSET_DIR = HELP_DIR / "assets"
TODAY = date.today().strftime("%B %-d, %Y") if os.name != "nt" else date.today().strftime("%B %#d, %Y")


NAVY = "#0B2341"
PANEL = "#12243A"
LINE = "#2B5A88"
CYAN = "#38BDF8"
GOLD = "#FFD166"
WHITE = "#F8FBFF"
INK = "#0F172A"
MUTED = "#64748B"
LIGHT_BG = "#F7FBFF"
LIGHT_LINE = "#CFE1F4"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        r"C:\Windows\Fonts\segoeuib.ttf" if bold else r"C:\Windows\Fonts\segoeui.ttf",
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def rounded(draw: ImageDraw.ImageDraw, box, radius=16, fill="#ffffff", outline=None, width=1):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def text(draw, xy, value, size=18, fill=INK, bold=False, anchor=None):
    draw.text(xy, value, font=font(size, bold), fill=fill, anchor=anchor)


def fit_text(draw, value, max_width, size=18, bold=False):
    current = value
    f = font(size, bold)
    if draw.textbbox((0, 0), current, font=f)[2] <= max_width:
        return current
    while len(current) > 3 and draw.textbbox((0, 0), current + "...", font=f)[2] > max_width:
        current = current[:-1]
    return current + "..."


def save(img: Image.Image, name: str):
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    HELP_ASSET_DIR.mkdir(parents=True, exist_ok=True)
    p = ASSET_DIR / name
    img.save(p)
    img.save(HELP_ASSET_DIR / name)
    return p


def draw_grid(draw, x, y, w, h, headers, rows):
    rounded(draw, (x, y, x + w, y + h), 14, fill=PANEL, outline=LINE, width=2)
    table_x, table_y = x + 18, y + 54
    table_w, table_h = w - 36, h - 74
    draw.rectangle((table_x, table_y, table_x + table_w, table_y + table_h), fill="#FFFFFF", outline="#D7DEE8")
    col_widths = [0.22, 0.28, 0.24, 0.13, 0.13]
    if len(headers) == 6:
        col_widths = [0.20, 0.22, 0.24, 0.11, 0.13, 0.10]
    cx = table_x
    for i, header in enumerate(headers):
        cw = int(table_w * col_widths[i])
        draw.rectangle((cx, table_y, cx + cw, table_y + 26), fill="#EEF4FA", outline="#D7DEE8")
        text(draw, (cx + 8, table_y + 6), header, 13, fill=INK)
        cx += cw
    row_y = table_y + 26
    row_h = 30
    for row in rows:
        cx = table_x
        for i, val in enumerate(row):
            cw = int(table_w * col_widths[i])
            draw.rectangle((cx, row_y, cx + cw, row_y + row_h), fill="#FFFFFF", outline="#E5E7EB")
            text(draw, (cx + 8, row_y + 7), fit_text(draw, str(val), cw - 14, 14), 14, fill=INK)
            cx += cw
        row_y += row_h
        if row_y + row_h > table_y + table_h:
            break


def dashboard_overview():
    img = Image.new("RGB", (1280, 760), "#08111F")
    d = ImageDraw.Draw(img)
    rounded(d, (28, 22, 1252, 118), 22, fill=NAVY)
    rounded(d, (54, 40, 112, 98), 12, fill="#0A1930", outline=LINE)
    d.rectangle((68, 80, 75, 92), fill=CYAN)
    d.rectangle((80, 72, 87, 92), fill=GOLD)
    d.rectangle((92, 62, 99, 92), fill=CYAN)
    d.line((68, 73, 78, 67, 88, 69, 100, 54), fill=CYAN, width=2)
    d.ellipse((96, 50, 104, 58), fill=GOLD)
    text(d, (134, 46), "Website Analytics", 28, WHITE, True)
    text(d, (136, 81), "Clear GA4 reporting", 16, WHITE)
    text(d, (430, 42), "Website property", 14, WHITE, True)
    d.rectangle((430, 63, 590, 96), fill="#FFFFFF")
    text(d, (440, 72), "Example Site", 14, INK)
    text(d, (610, 42), "Date range", 14, WHITE, True)
    d.rectangle((610, 63, 750, 96), fill="#FFFFFF")
    text(d, (620, 72), "Last 7 days", 14, INK)
    d.rectangle((770, 76, 780, 86), outline=CYAN, fill="#102B4D")
    d.line((772, 81, 775, 84, 780, 77), fill=CYAN, width=2)
    text(d, (788, 72), "Auto 60s", 14, WHITE)
    rounded(d, (855, 36, 1225, 104), 18, fill="#102B4D", outline=LINE)
    text(d, (884, 52), "Status", 13, WHITE)
    text(d, (884, 75), "Connected", 22, WHITE, True)
    for i, label in enumerate(["Update", "Settings", "Help"]):
        bx = 1015 + i * 70
        bw = 62 if label != "Settings" else 74
        rounded(d, (bx, 54, bx + bw, 88), 6, fill="#FFFFFF", outline="#D8DEE8")
        text(d, (bx + bw / 2, 65), label, 12, INK, anchor="ma")

    rounded(d, (16, 150, 618, 380), 14, fill=PANEL, outline=LINE, width=2)
    text(d, (36, 170), "Country, region/state, city, users, and sessions", 15, GOLD)
    draw_grid(d, 16, 150, 602, 230, ["Country", "Region / State", "City", "Users", "Sessions"], [
        ("Australia", "Western Australia", "Perth", 1, 1),
        ("Brazil", "Federal District", "(city not set)", 1, 1),
        ("India", "Maharashtra", "Mumbai", 3, 4),
        ("United States", "Georgia", "Dublin", 7, 12),
    ])

    rounded(d, (16, 398, 618, 736), 14, fill=PANEL, outline=LINE, width=2)
    text(d, (36, 418), "Top pages and downloads for Last 7 days", 15, GOLD)
    bar_items = [
        ("Home - Page view", 58),
        ("/downloads.html - Download: file_download", 48),
        ("Home - session_start", 43),
        ("/downloads.html - Page view", 38),
        ("Home - user_engagement", 31),
    ]
    maxv = max(v for _, v in bar_items)
    for i, (label, value) in enumerate(bar_items):
        yy = 460 + i * 42
        text(d, (36, yy), fit_text(d, label, 300, 14), 14, WHITE)
        rounded(d, (335, yy - 3, 335 + int(250 * value / maxv), yy + 16), 5, fill=CYAN)
        text(d, (600, yy - 2), str(value), 14, GOLD, anchor="ra")

    rounded(d, (650, 150, 1248, 380), 14, fill=PANEL, outline=LINE, width=2)
    text(d, (670, 174), "Users over the last 7 days (7/16/2026 - 7/22/2026)", 18, WHITE, True)
    chart_x, chart_y, chart_w, chart_h = 705, 225, 500, 115
    for i in range(5):
        y = chart_y + chart_h - int(i * chart_h / 4)
        d.line((chart_x, y, chart_x + chart_w, y), fill="#2A5278", width=1)
        text(d, (chart_x - 24, y - 7), str(i * 30), 12, "#B8C7DB")
    vals = [4, 8, 46, 86, 37, 30, 9]
    for i, v in enumerate(vals):
        bw = 44
        gap = 28
        bx = chart_x + 36 + i * (bw + gap)
        bh = int(chart_h * v / 120)
        rounded(d, (bx, chart_y + chart_h - bh, bx + bw, chart_y + chart_h), 6, fill=CYAN)
        text(d, (bx + bw / 2, chart_y + chart_h + 12), f"7/{16+i}", 12, "#B8C7DB", anchor="ma")
    rounded(d, (860, 210, 955, 255), 7, fill="#071323", outline=CYAN)
    text(d, (872, 220), "7/19", 12, WHITE)
    text(d, (872, 238), "Users: 86", 12, WHITE, True)

    rounded(d, (650, 398, 1248, 736), 14, fill=PANEL, outline=LINE, width=2)
    text(d, (682, 430), "Users today", 14, "#B8C7DB")
    text(d, (682, 466), "3", 34, GOLD, True)
    text(d, (840, 430), "Active now", 14, "#B8C7DB")
    text(d, (840, 466), "1", 34, GOLD, True)
    text(d, (1000, 430), "Last Location", 14, "#B8C7DB")
    text(d, (1000, 466), "Dublin, United States", 20, GOLD, True)
    text(d, (682, 550), "Current location: Dublin, United States", 17, WHITE)
    text(d, (682, 580), "Last activity: on 7/22/2026 at 8:42:15 AM", 17, WHITE)
    return save(img, "dashboard-overview-sanitized.png")


def summary_cards():
    img = Image.new("RGB", (840, 250), PANEL)
    d = ImageDraw.Draw(img)
    titles = [
        ("Users today", "3"),
        ("Users for\nLast 7 days", "132"),
        ("Views today", "12"),
        ("Downloads\nLast 7 days", "226"),
        ("Active now", "1"),
        ("Last Location", "Dublin,\nUnited States"),
        ("Top page", "Home page"),
        ("Views last 30 min", "1"),
        ("Events/session", "1.6"),
        ("Scrolled", "42"),
    ]
    for i, (title, value) in enumerate(titles):
        row = i // 5
        col = i % 5
        x = 20 + col * 160
        y = 18 + row * 112
        rounded(d, (x, y, x + 145, y + 94), 12, fill="#102B4D", outline=LINE)
        for j, line in enumerate(title.split("\n")):
            text(d, (x + 14, y + 14 + j * 16), line, 14, "#B8C7DB")
        val_size = 18 if "\n" in value else 25
        for j, line in enumerate(value.split("\n")):
            text(d, (x + 14, y + 56 + j * 22), line, val_size, GOLD, True)
    return save(img, "summary-cards-sanitized.png")


def users_graph():
    img = Image.new("RGB", (720, 330), PANEL)
    d = ImageDraw.Draw(img)
    rounded(d, (8, 8, 712, 322), 12, fill=PANEL, outline=CYAN, width=4)
    text(d, (28, 28), "Users over the last 7 days (7/16/2026 - 7/22/2026)", 18, WHITE, True)
    chart_x, chart_y, chart_w, chart_h = 90, 80, 560, 190
    for i, label in enumerate([0, 30, 60, 90, 120]):
        y = chart_y + chart_h - int(i * chart_h / 4)
        d.line((chart_x, y, chart_x + chart_w, y), fill="#31597D")
        text(d, (chart_x - 30, y - 7), str(label), 12, "#C7D6EA")
    vals = [4, 8, 46, 86, 37, 30, 9]
    for i, v in enumerate(vals):
        bw = 48
        gap = 30
        bx = chart_x + i * (bw + gap) + 25
        bh = int(chart_h * v / 120)
        rounded(d, (bx, chart_y + chart_h - bh, bx + bw, chart_y + chart_h), 7, fill=CYAN)
        text(d, (bx + bw / 2, chart_y + chart_h + 15), f"7/{16+i}", 12, "#C7D6EA", anchor="ma")
    rounded(d, (310, 93, 420, 145), 7, fill="#071323", outline=CYAN, width=1)
    text(d, (324, 105), "7/19", 13, WHITE)
    text(d, (324, 127), "Users: 86", 13, WHITE, True)
    return save(img, "users-graph-crop-sanitized.png")


def location_grid():
    img = Image.new("RGB", (760, 255), PANEL)
    d = ImageDraw.Draw(img)
    text(d, (20, 20), "Country, region/state, city, users, and sessions", 16, GOLD)
    draw_grid(d, 0, 0, 760, 255, ["Country", "Region / State", "City", "Users", "Sessions"], [
        ("Australia", "Western Australia", "Perth", 1, 1),
        ("Brazil", "Federal District", "(city not set)", 1, 1),
        ("India", "Maharashtra", "Mumbai", 3, 4),
        ("United States", "Georgia", "Dublin", 7, 12),
        ("United States", "Arizona", "Phoenix", 1, 1),
    ])
    return save(img, "locations-grid-crop-sanitized.png")


def pages_chart():
    img = Image.new("RGB", (760, 330), PANEL)
    d = ImageDraw.Draw(img)
    rounded(d, (6, 6, 754, 324), 14, fill=PANEL, outline=LINE, width=2)
    text(d, (24, 24), "Top pages and downloads for Last 7 days", 16, GOLD)
    items = [
        ("Home - Page view", 58),
        ("/downloads.html - Download: file_download", 48),
        ("Home - session_start", 43),
        ("/downloads.html - Page view", 38),
        ("Home - user_engagement", 31),
        ("/downloads.html - scroll", 22),
    ]
    maxv = max(v for _, v in items)
    for i, (label, value) in enumerate(items):
        yy = 72 + i * 38
        text(d, (26, yy), fit_text(d, label, 330, 15), 15, WHITE)
        rounded(d, (370, yy - 4, 370 + int(285 * value / maxv), yy + 17), 5, fill=CYAN)
        text(d, (705, yy - 2), str(value), 15, GOLD, anchor="ra")
    return save(img, "pages-downloads-chart-sanitized.png")


def realtime_panel():
    img = Image.new("RGB", (720, 210), PANEL)
    d = ImageDraw.Draw(img)
    rounded(d, (6, 6, 714, 204), 14, fill=PANEL, outline=LINE, width=2)
    text(d, (28, 42), "Current active visitors", 16, "#B8C7DB")
    text(d, (28, 86), "1", 44, GOLD, True)
    text(d, (250, 66), "Current location: Dublin, United States", 18, WHITE)
    text(d, (250, 98), "Last activity: on 7/22/2026 at 8:42:15 AM", 18, WHITE)
    return save(img, "realtime-panel-crop-sanitized.png")


def settings_card():
    img = Image.new("RGB", (760, 430), LIGHT_BG)
    d = ImageDraw.Draw(img)
    rounded(d, (20, 20, 740, 410), 18, fill="#FFFFFF", outline=LIGHT_LINE, width=2)
    text(d, (48, 48), "Settings and Google Sign-in", 28, NAVY, True)
    text(d, (48, 88), "Use your own Google desktop OAuth client. Secrets are hidden by default.", 15, MUTED)
    fields = [
        ("Default website", "Example Site"),
        ("Default date range", "Last 7 days"),
        ("OAuth client ID", "123456789012-abc...apps.googleusercontent.com"),
        ("OAuth client secret", "••••••••••••••••"),
        ("Saved refresh token", "Encrypted with Windows DPAPI"),
        ("Redirect URI", "http://127.0.0.1:53682/oauth2redirect"),
    ]
    y = 125
    for label, value in fields:
        text(d, (54, y), label, 14, NAVY, True)
        rounded(d, (210, y - 8, 710, y + 27), 7, fill="#F8FBFF", outline=LIGHT_LINE)
        text(d, (222, y), value, 15, INK)
        y += 47
    return save(img, "settings-redacted-clean.png")


def properties_card():
    img = Image.new("RGB", (760, 350), LIGHT_BG)
    d = ImageDraw.Draw(img)
    rounded(d, (20, 20, 740, 330), 18, fill="#FFFFFF", outline=LIGHT_LINE, width=2)
    text(d, (48, 50), "Managing Properties", 28, NAVY, True)
    text(d, (48, 88), "Use generic names in examples; enter your own GA4 numeric property IDs.", 15, MUTED)
    rows = [
        ("Example Site A", "123456789"),
        ("Example Site B", "234567890"),
        ("Example Site C", "345678901"),
    ]
    y = 130
    for name, pid in rows:
        rounded(d, (55, y, 705, y + 48), 10, fill="#F8FBFF", outline=LIGHT_LINE)
        text(d, (75, y + 14), name, 16, INK, True)
        text(d, (575, y + 14), pid, 16, MUTED)
        y += 62
    return save(img, "properties-redacted-clean.png")


def diagnostics_card():
    img = Image.new("RGB", (760, 310), LIGHT_BG)
    d = ImageDraw.Draw(img)
    rounded(d, (20, 20, 740, 290), 18, fill="#FFFFFF", outline=LIGHT_LINE, width=2)
    text(d, (48, 50), "Diagnostics", 28, NAVY, True)
    items = [
        ("Authentication", "Connected"),
        ("Selected property", "Example Site"),
        ("Last GA4 request", "Successful"),
        ("Startup token refresh", "Enabled"),
    ]
    y = 105
    for label, value in items:
        text(d, (60, y), label, 16, INK, True)
        rounded(d, (420, y - 8, 680, y + 27), 8, fill="#EAF8FF", outline="#9BDCF8")
        text(d, (440, y), value, 16, NAVY)
        y += 42
    return save(img, "diagnostics-clean.png")


def privacy_card():
    img = Image.new("RGB", (760, 360), LIGHT_BG)
    d = ImageDraw.Draw(img)
    rounded(d, (20, 20, 740, 340), 18, fill="#FFFFFF", outline=LIGHT_LINE, width=2)
    text(d, (48, 48), "Privacy and Storage", 28, NAVY, True)
    rows = [
        ("Stored locally", "Settings, property definitions, saved grid widths"),
        ("Stored encrypted", "Google refresh token via Windows DPAPI"),
        ("Not stored", "Fetched GA4 report rows and analytics results"),
        ("Ignored by Git", "databases/*.sqlite3"),
        ("Tracked by Git", "databases/WebsiteAnalytics.schema.sql"),
    ]
    y = 105
    for label, value in rows:
        rounded(d, (55, y, 705, y + 38), 8, fill="#F8FBFF", outline=LIGHT_LINE)
        text(d, (75, y + 9), label, 15, NAVY, True)
        text(d, (235, y + 9), value, 15, INK)
        y += 48
    return save(img, "privacy-storage-clean.png")


def create_assets():
    for old in ASSET_DIR.glob("*.png"):
        old.unlink()
    assets = [
        dashboard_overview(),
        summary_cards(),
        users_graph(),
        location_grid(),
        pages_chart(),
        realtime_panel(),
        settings_card(),
        properties_card(),
        diagnostics_card(),
        privacy_card(),
    ]
    return assets


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.replace("#", ""))
    tc_pr.append(shd)


def set_cell_margins(table):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for m in ["top", "start", "bottom", "end"]:
        node = tbl_cell_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), "120")
        node.set(qn("w:type"), "dxa")


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def style_doc(doc, title):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    for style_name, size, color in [
        ("Heading 1", 18, NAVY),
        ("Heading 2", 14, "#1974DF"),
        ("Heading 3", 12, "#1F4D78"),
    ]:
        st = styles[style_name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(6)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor.from_string(NAVY.replace("#", ""))
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"Website Analytics GA4 Dashboard • Last changed {TODAY}")
    r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor.from_string("64748B")
    doc.add_paragraph("")
    callout = doc.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_margins(callout)
    cell = callout.cell(0, 0)
    set_cell_shading(cell, "#EAF8FF")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = "This guide uses sanitized examples. Replace the sample site names, property IDs, OAuth client ID, and OAuth secret with your own values."
    doc.add_paragraph("")
    doc.add_heading("Table of Contents", level=1)
    add_toc(doc.add_paragraph())
    doc.add_page_break()


def add_figure(doc, image_name, caption, width=6.4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(ASSET_DIR / image_name), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.italic = True
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor.from_string("64748B")


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def user_guide():
    doc = Document()
    style_doc(doc, "Website Analytics User Guide")
    doc.add_heading("What the Dashboard Does", level=1)
    doc.add_paragraph(
        "Website Analytics is a desktop dashboard for GA4 data. It opens, reconnects to Google when possible, retrieves live report data into memory, and presents the information in a simpler dashboard than the standard GA4 web console."
    )
    add_figure(doc, "dashboard-overview-sanitized.png", "The current dashboard layout with sanitized sample data.")
    doc.add_heading("Daily Workflow", level=1)
    add_numbered(doc, [
        "Choose the website property. Use All websites only when you want the app to combine enabled properties.",
        "Choose the date range. The cards, location grid, users chart, and top pages/downloads chart follow this selection.",
        "Leave Auto 60 sec checked if you want the app to keep refreshing while it runs.",
        "Use Update when you want an immediate refresh.",
    ])
    doc.add_heading("Summary Cards", level=1)
    doc.add_paragraph("The card strip gives a fast read on the most useful numbers without scrolling through GA4.")
    add_figure(doc, "summary-cards-sanitized.png", "Ten dashboard cards: users, views, downloads, realtime status, location, top page, events/session, and scrolled users.", 6.2)
    add_bullets(doc, [
        "Users today counts users for the current day.",
        "Users for the selected range follows the Date range selector.",
        "Downloads follows the selected range and counts download events when GA4 receives them.",
        "Active now and Views last 30 min use GA4 realtime reporting.",
        "Events/session and Scrolled are GA4 engagement indicators.",
    ])
    doc.add_heading("Locations and Pages", level=1)
    doc.add_paragraph("The location grid groups users and sessions by country, region/state, and city. The pages/downloads chart shows the most-used pages and actions for the selected range.")
    add_figure(doc, "locations-grid-crop-sanitized.png", "A cropped location grid. The examples are sanitized and do not show real property data.", 6.2)
    add_figure(doc, "pages-downloads-chart-sanitized.png", "Top pages and downloads shown as a readable bar chart instead of a crowded table.", 6.2)
    doc.add_heading("Users Graph and Tooltips", level=1)
    doc.add_paragraph("The users chart shows users over the selected date range. Hover over a bar to see the exact date and value.")
    add_figure(doc, "users-graph-crop-sanitized.png", "A cropped users chart with a compact tooltip.", 6.0)
    doc.add_heading("Realtime Activity", level=1)
    doc.add_paragraph("The realtime panel answers the question: what is happening right now? It shows current active visitors, current location when GA4 returns it, and the most recent activity time.")
    add_figure(doc, "realtime-panel-crop-sanitized.png", "Realtime activity panel with sanitized sample location information.", 6.0)
    doc.add_heading("Settings and Privacy", level=1)
    doc.add_paragraph("Settings hold dashboard defaults, the Google OAuth desktop client values, and the encrypted refresh token. Report rows are fetched into memory and discarded when the app closes.")
    add_figure(doc, "settings-redacted-clean.png", "Settings are documented with dummy OAuth values and hidden secrets.", 6.1)
    add_figure(doc, "privacy-storage-clean.png", "Privacy and storage summary using a brighter documentation style.", 6.1)
    doc.add_heading("Managing Properties", level=1)
    doc.add_paragraph("Add one row for each website you want to report on. Use the numeric GA4 property ID, not a measurement ID.")
    add_figure(doc, "properties-redacted-clean.png", "Property examples use dummy names and dummy numeric IDs.", 6.1)
    doc.add_heading("Troubleshooting", level=1)
    add_bullets(doc, [
        "If Google sign-in appears again, the saved refresh token may be missing, revoked, or tied to a different Windows user.",
        "If GA4 says permission is denied, confirm the signed-in Google account has access to the GA4 property.",
        "If a card shows zero, compare the selected date range with GA4 and remember that realtime and standard reports are different GA4 APIs.",
        "If a download count seems low, confirm the website is emitting the expected GA4 download event.",
    ])
    doc.save(GUIDES_DIR / "WebsiteAnalytics_User_Guide.docx")


def engineering_guide():
    doc = Document()
    style_doc(doc, "Website Analytics Engineering Guide")
    doc.add_heading("Architecture Overview", level=1)
    doc.add_paragraph(
        "The application is a Delphi FMX desktop dashboard. Forms remain editable in the IDE. Runtime code retrieves GA4 data, maps responses into memory models, and updates existing designer controls."
    )
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_margins(tbl)
    for i, h in enumerate(["Area", "Primary unit", "Purpose"]):
        tbl.cell(0, i).text = h
        set_cell_shading(tbl.cell(0, i), "#E8EEF5")
    rows = [
        ("Dashboard form", "WebsiteAnalytics.MainForm.pas/.fmx", "Hero, summary cards, chart paint boxes, grids, realtime panel, update timers."),
        ("GA4 access", "WebsiteAnalytics.GA4DataModule.pas", "Builds GA4 Data API and Realtime API requests and parses responses."),
        ("Authentication", "WebsiteAnalytics.AuthenticationDataModule.pas", "Desktop OAuth, loopback redirect, refresh-token persistence."),
        ("Settings", "WebsiteAnalytics.SettingsDataModule.pas", "Portable SQLite settings database and schema."),
        ("Models", "WebsiteAnalytics.Models.pas", "Snapshots, KPI summaries, trend points, location/content rows."),
    ]
    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_heading("Current Dashboard Surface", level=1)
    add_figure(doc, "dashboard-overview-sanitized.png", "Current surface area with sanitized data.")
    doc.add_heading("GA4 Reports Used", level=1)
    add_bullets(doc, [
        "Overview/KPI requests include users, sessions, screen/page views, engagement rate, event count, events per session, and scrolled users.",
        "Trend requests normalize missing dates so the chart can show a complete selected range.",
        "Geography requests group country, region, city, users, sessions, and engagement information.",
        "Content requests group page title, path, action/event, users, and event/download counts.",
        "Realtime requests read active users, realtime views, and the current/last activity location when GA4 returns those dimensions.",
    ])
    doc.add_heading("Authentication and Token Storage", level=1)
    doc.add_paragraph(
        "The app uses a Google desktop OAuth client. On first sign-in the loopback listener receives an authorization code, exchanges it for tokens, and stores the refresh token encrypted for the current Windows user. Startup attempts a silent token refresh before showing any browser login screens."
    )
    add_figure(doc, "settings-redacted-clean.png", "OAuth settings are shown with dummy values and a hidden client secret.", 6.1)
    doc.add_heading("Portable SQLite Settings", level=1)
    doc.add_paragraph(
        "The SQLite database stores application settings, property definitions, saved grid widths, OAuth setup values, and the encrypted refresh token. Report data remains memory-only."
    )
    add_figure(doc, "privacy-storage-clean.png", "Storage responsibilities and Git tracking policy.", 6.1)
    doc.add_heading("Designer/Editability Rules", level=1)
    add_bullets(doc, [
        "FMX forms should remain editable in RAD Studio.",
        "Do not add helper code or runtime UI layout code without developer approval.",
        "Prefer Object Inspector settings for control placement and sizing whenever practical.",
        "When runtime drawing is necessary for charts or tooltips, keep it limited to painting/data display, not form construction.",
    ])
    doc.add_heading("Build and Validation", level=1)
    add_numbered(doc, [
        "Create a pre-change backup on D: before modifying project files.",
        "Build Win64 and Win32 Debug with RAD Studio 37.0 MSBuild.",
        "Smoke-test startup to catch FMX read errors.",
        "Review Git status and stage only intended files.",
        "Commit and push both the private and public repositories when the change is complete.",
    ])
    doc.add_heading("Public Repository Hygiene", level=1)
    add_bullets(doc, [
        "Never commit a real SQLite database, OAuth client secret, refresh token, or personal property IDs.",
        "Keep databases/WebsiteAnalytics.schema.sql tracked.",
        "Use dummy property IDs and redacted OAuth examples in documentation.",
        "Keep generated binaries and local build outputs ignored unless an explicit release artifact is being created.",
    ])
    add_figure(doc, "properties-redacted-clean.png", "Sanitized property examples suitable for public documentation.", 6.1)
    doc.save(GUIDES_DIR / "WebsiteAnalytics_Engineering_Guide.docx")


def topic_html(title, image, caption, body):
    paragraphs = "\n".join(f"<p>{p}</p>" for p in body.get("p", []))
    bullets = ""
    if body.get("bullets"):
        bullets = "<ul>" + "".join(f"<li>{b}</li>" for b in body["bullets"]) + "</ul>"
    sections = ""
    for heading, items in body.get("sections", []):
        sections += f"<h2>{heading}</h2><ul>" + "".join(f"<li>{i}</li>" for i in items) + "</ul>"
    return f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <title>{title}</title>
  <link rel="stylesheet" href="../style.css">
</head>
<body class="topic-body">
<main class="topic">
<h1>{title}</h1>
<figure><img src="../assets/{image}" alt="{title}"><figcaption>{caption}</figcaption></figure>
{paragraphs}
{bullets}
{sections}
</main>
</body>
</html>
"""


def update_help():
    css = """:root {
  --navy: #0b2341;
  --blue: #1974df;
  --accent: #38bdf8;
  --gold: #b77900;
  --ink: #0f172a;
  --muted: #64748b;
  --page: #f7fbff;
  --card: #ffffff;
  --line: #cfe1f4;
}
body { margin: 0; font-family: Segoe UI, Arial, sans-serif; color: var(--ink); background: var(--page); }
.help-shell { display: grid; grid-template-columns: 330px 1fr; height: 100vh; }
aside { background: linear-gradient(180deg, #ffffff 0%, #eaf8ff 100%); padding: 18px; overflow: auto; border-right: 1px solid var(--line); }
aside h1 { margin: 0 0 4px; font-size: 24px; color: var(--navy); }
.help-subtitle { margin: 0 0 14px; color: var(--muted); font-size: 13px; }
.shell-search input { width: 100%; box-sizing: border-box; padding: 10px; border-radius: 10px; border: 1px solid var(--line); background: #fff; }
#searchResults a, #searchResults span { display: block; color: var(--blue); padding: 6px 2px; }
.shell-links { display: flex; gap: 8px; margin: 12px 0; flex-wrap: wrap; }
.shell-links a { background: #ffffff; border: 1px solid var(--line); border-radius: 10px; padding: 7px 9px; color: var(--navy); text-decoration: none; }
nav ul { list-style: none; padding: 0; margin: 12px 0 0; }
nav li { margin: 7px 0; }
nav a { display: block; padding: 10px 11px; border-radius: 12px; background: #ffffff; border: 1px solid var(--line); color: var(--navy); text-decoration: none; font-weight: 650; box-shadow: 0 2px 8px rgba(15, 23, 42, 0.05); }
nav a.active, nav a:hover { background: var(--accent); color: #06111f; }
.topic-frame { background: #eaf8ff; padding: 14px; }
iframe { width: 100%; height: calc(100vh - 28px); border: 1px solid var(--line); border-radius: 16px; background: white; }
.topic-body { background: var(--page); color: var(--ink); }
.topic { max-width: 1040px; margin: 0 auto; padding: 32px 38px 60px; }
.topic h1 { color: var(--navy); border-bottom: 4px solid var(--accent); padding-bottom: 9px; margin-top: 0; font-size: 32px; }
.topic h2 { color: var(--blue); margin-top: 28px; font-size: 22px; }
.topic p, .topic li { font-size: 16px; line-height: 1.58; }
.topic li { margin-bottom: 8px; }
figure { margin: 18px 0 26px; border: 1px solid var(--line); border-radius: 16px; padding: 12px; background: var(--card); box-shadow: 0 10px 28px rgba(15, 23, 42, 0.10); }
figure img { display: block; max-width: 100%; height: auto; border-radius: 10px; margin: 0 auto; }
figcaption { color: var(--muted); font-size: 13px; padding-top: 8px; text-align: center; }
.note { border-left: 5px solid var(--accent); background: #ffffff; padding: 12px 14px; border-radius: 10px; }
.print-topic { page-break-after: always; }
@media print { .topic { max-width: none; padding: 24px; } figure { box-shadow: none; } }
"""
    (HELP_DIR / "style.css").write_text(css, encoding="utf-8")
    topics = {
        "overview.html": ("Overview", "dashboard-overview-sanitized.png", "The current dashboard uses cropped, sanitized examples in this help system.", {
            "p": ["Website Analytics presents GA4 data in a focused desktop dashboard. The main screen is organized around locations, pages/downloads, trend activity, realtime activity, and summary cards."],
            "bullets": ["Report rows are held in memory only.", "Settings are stored locally in SQLite.", "The saved Google refresh token is encrypted with Windows DPAPI."],
        }),
        "quick-start.html": ("Quick Start", "settings-redacted-clean.png", "Settings with dummy OAuth values and a hidden client secret.", {
            "sections": [
                ("First run", ["Create a Google desktop OAuth client.", "Enter the client ID and secret in Settings.", "Confirm the redirect URI matches your OAuth desktop client.", "Connect Google once, then save settings."]),
                ("Normal use", ["Choose a property and date range.", "Leave Auto 60 sec enabled for a live wallboard-style dashboard.", "Use Update for an immediate refresh."]),
            ],
        }),
        "dashboard-tour.html": ("Dashboard Tour", "dashboard-overview-sanitized.png", "The dashboard surface with sanitized example data.", {
            "p": ["The dashboard has a compact hero section, summary cards, a location grid, a top pages/downloads chart, a users trend graph, and a realtime panel."],
        }),
        "website-date-filters.html": ("Website and Date Filters", "dashboard-overview-sanitized.png", "The hero area contains the property selector, date range selector, and auto-update checkbox.", {
            "bullets": ["Website property controls which GA4 property is queried.", "Date range controls standard reports, charts, and summary cards.", "All websites combines enabled properties.", "Realtime values come from GA4 realtime endpoints and may not exactly match standard reports at the same instant."],
        }),
        "trend-graph.html": ("Users Graph", "users-graph-crop-sanitized.png", "A cropped users graph with exact hover tooltip.", {
            "bullets": ["The graph follows the selected date range.", "The axis uses whole numbers.", "Hover over a bar to read the exact value."],
        }),
        "locations.html": ("Visitor Locations", "locations-grid-crop-sanitized.png", "Locations grouped by country, region/state, and city.", {
            "bullets": ["The grid is sorted alphabetically for easier lookup.", "United States rows use city and state where GA4 returns them.", "International rows use city and country/region where available.", "(not set) means GA4 did not return that dimension."],
        }),
        "pages-actions.html": ("Pages, Downloads, and Actions", "pages-downloads-chart-sanitized.png", "Top pages and downloads shown as a readable chart.", {
            "bullets": ["Home page is displayed when GA4 reports the path /.", "Download events appear when the website emits file_download or download-style events.", "The chart follows the selected date range."],
        }),
        "realtime.html": ("Realtime Activity", "realtime-panel-crop-sanitized.png", "Realtime panel showing active visitor count and last activity.", {
            "bullets": ["Active now is the current realtime active user count.", "Current location appears when GA4 returns city/country or city/state.", "Last activity is reset when the selected property changes."],
        }),
        "properties.html": ("Managing Properties", "properties-redacted-clean.png", "Public documentation uses dummy property names and IDs.", {
            "bullets": ["Use the numeric GA4 property ID.", "Do not use the Measurement ID that begins with G-.", "Add more properties as you add more websites."],
        }),
        "settings-and-google.html": ("Settings and Google Sign-in", "settings-redacted-clean.png", "OAuth examples are redacted and safe for public documentation.", {
            "bullets": ["The client secret is hidden by default in the UI.", "The refresh token is encrypted locally.", "The app silently reconnects on startup when the saved token is valid.", "Google login appears only when the token is missing, revoked, or cannot be refreshed."],
        }),
        "long-running-dashboard.html": ("Running the Dashboard All Day", "summary-cards-sanitized.png", "Summary cards are designed for at-a-glance monitoring.", {
            "bullets": ["Keep the app open and Auto 60 sec enabled.", "Use the status bar for the latest update message.", "If authentication expires, open Settings and reconnect Google."],
        }),
        "diagnostics.html": ("Diagnostics", "diagnostics-clean.png", "Diagnostics show connection and request status without exposing secrets.", {
            "bullets": ["Use Diagnostics when Update fails.", "Check authentication state and selected property.", "Look at the last GA4 request message for permission or API errors."],
        }),
        "troubleshooting.html": ("Troubleshooting", "diagnostics-clean.png", "Diagnostics are the first stop for GA4 errors.", {
            "sections": [
                ("Common issues", ["403 permission errors usually mean the signed-in account cannot read that GA4 property.", "Zero data can be normal for quiet sites or dates with no GA4 traffic.", "Realtime and standard GA4 reports can update on different schedules."]),
                ("What to check", ["OAuth client ID and secret.", "GA4 property ID.", "Google Analytics Data API enabled in Google Cloud.", "Tester/verification status for the OAuth consent screen."]),
            ],
        }),
        "privacy-storage.html": ("Privacy and Storage", "privacy-storage-clean.png", "Storage rules shown in a brighter, more readable help style.", {
            "bullets": ["Settings and property definitions are local.", "Refresh tokens are encrypted for the current Windows user.", "GA4 report rows are not saved.", "The repository should track the schema, not a real SQLite database."],
        }),
    }
    for filename, args in topics.items():
        (HELP_DIR / "topics" / filename).write_text(topic_html(*args), encoding="utf-8")
    manual = ["<!doctype html><html lang=\"en\"><head><meta charset=\"utf-8\"><title>Website Analytics Help Manual</title><link rel=\"stylesheet\" href=\"style.css\"></head><body class=\"topic-body\">"]
    for filename in topics:
        content = (HELP_DIR / "topics" / filename).read_text(encoding="utf-8")
        body = re.search(r"<main class=\"topic\">(.*?)</main>", content, re.S).group(1)
        manual.append(f"<section class=\"topic print-topic\">{body}</section>")
    manual.append("</body></html>")
    (HELP_DIR / "manual-print.html").write_text("\n".join(manual), encoding="utf-8")
    metadata = {
        "name": "Website Analytics Help",
        "outputPath": "help",
        "lastChanged": TODAY,
        "topics": json.loads((HELP_DIR / "toc.json").read_text(encoding="utf-8")),
    }
    (HELP_DIR / "WebsiteAnalytics Help.hdc.json").write_text(json.dumps(metadata, indent=2), encoding="utf-8")
    (HELP_DIR / "metadata.json").write_text(json.dumps({"lastChanged": TODAY, "generator": "tools/refresh_websiteanalytics_docs.py"}, indent=2), encoding="utf-8")
    search = []
    for filename, (title, _, _, body) in topics.items():
        text_bits = [title] + body.get("p", []) + body.get("bullets", [])
        for section_title, items in body.get("sections", []):
            text_bits.append(section_title)
            text_bits.extend(items)
        search.append({"title": title, "url": f"topics/{filename}", "text": " ".join(text_bits)})
    (HELP_DIR / "search-index.json").write_text(json.dumps(search, indent=2), encoding="utf-8")
    nav = "".join(
        f'<li><a href="{item["url"]}" target="topicFrame" data-topic-url="{item["url"]}" '
        f'onclick="setActiveTopic(this.getAttribute(\'data-topic-url\'))">{item["title"]}</a></li>'
        for item in json.loads((HELP_DIR / "toc.json").read_text(encoding="utf-8"))
    )
    index_html = f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <title>Website Analytics Help</title>
  <link rel="stylesheet" href="style.css">
</head>
<body class="help-shell">
  <aside>
    <h1>Website Analytics Help</h1>
    <p class="help-subtitle">Friendly local help for the GA4 desktop dashboard.</p>
    <div class="shell-search">
      <input id="searchBox" placeholder="Search help" oninput="runSearch()">
      <div id="searchResults"></div>
    </div>
    <div class="shell-links">
      <a href="manual-print.html" target="topicFrame">Print view</a>
      <a href="keyword-index.html" target="topicFrame">Index</a>
      <a href="glossary.html" target="topicFrame">Glossary</a>
    </div>
    <nav><ul>{nav}</ul></nav>
  </aside>
  <section class="topic-frame">
    <iframe id="topicFrame" name="topicFrame" src="topics/overview.html" title="Help topic"></iframe>
  </section>
  <script>
    var searchItems = {json.dumps(search)};
    function cleanTopicUrl(url) {{
      url = String(url || '').split('#')[0].split('?')[0].replace(/\\\\/g, '/');
      var i = url.lastIndexOf('topics/');
      if (i >= 0) return url.substring(i);
      var p = url.lastIndexOf('/');
      if (p >= 0) return 'topics/' + url.substring(p + 1);
      return url;
    }}
    function clearActiveTopic() {{
      document.querySelectorAll('nav a.active').forEach(function(a) {{ a.classList.remove('active'); }});
    }}
    function setActiveTopic(url) {{
      var wanted = cleanTopicUrl(url);
      clearActiveTopic();
      document.querySelectorAll('nav a[data-topic-url]').forEach(function(a) {{
        if (cleanTopicUrl(a.getAttribute('href')) === wanted) {{
          a.classList.add('active');
          a.scrollIntoView({{block: 'nearest'}});
        }}
      }});
    }}
    function runSearch() {{
      var q = document.getElementById('searchBox').value.toLowerCase();
      var r = document.getElementById('searchResults');
      if (!q) {{ r.innerHTML = ''; return; }}
      var h = '';
      searchItems.filter(function(x) {{
        return (x.title + ' ' + x.text).toLowerCase().indexOf(q) >= 0;
      }}).slice(0, 20).forEach(function(x) {{
        h += '<a target="topicFrame" href="' + x.url + '" onclick="setActiveTopic(\\'' + x.url + '\\')">' + x.title + '</a>';
      }});
      r.innerHTML = h || '<span>No matches</span>';
    }}
    document.getElementById('topicFrame').addEventListener('load', function() {{
      try {{ setActiveTopic(this.contentWindow.location.href || this.contentWindow.location.pathname); }} catch(e) {{}}
    }});
    setActiveTopic('topics/overview.html');
  </script>
</body>
</html>
"""
    (HELP_DIR / "index.html").write_text(index_html, encoding="utf-8")


def main():
    create_assets()
    user_guide()
    engineering_guide()
    update_help()


if __name__ == "__main__":
    main()
