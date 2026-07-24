from __future__ import annotations

import subprocess
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
GUIDES_DIR = PROJECT_ROOT / "docs" / "guides"
PDF_DIR = PROJECT_ROOT / "docs" / "pdf"
ASSETS_DIR = PROJECT_ROOT / "help" / "assets"
DOCX_PATH = GUIDES_DIR / "WebsiteAnalytics_Public_How_To_Guide.docx"
PDF_PATH = PDF_DIR / "WebsiteAnalytics_Public_How_To_Guide.pdf"

LAST_CHANGED = date.today().strftime("%B %#d, %Y")
PUBLIC_REPO_URL = "https://github.com/tmartindub/WebsiteAnalytics-GA4-Dashboard.git"
EXAMPLE_PROJECT_FOLDER = r"C:\Projects\WebsiteAnalytics"
DEFAULT_REDIRECT_URI = "http://127.0.0.1:53682/oauth2redirect"
GA4_READONLY_SCOPE = "https://www.googleapis.com/auth/analytics.readonly"

ACCENT = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
TITLE_BLUE = RGBColor(11, 37, 69)
INK = RGBColor(20, 33, 50)
MUTED = RGBColor(90, 104, 120)
GOLD = RGBColor(122, 90, 0)
RED = RGBColor(155, 28, 28)
WHITE = RGBColor(255, 255, 255)
LIGHT_BLUE = "E8EEF5"
LIGHT_CYAN = "EAF8FF"
LIGHT_GOLD = "FFF4D6"
LIGHT_RED = "FDECEC"
TABLE_HEADER = "E8EEF5"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tbl_cell_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360, indent_dxa=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    table.autofit = False
    set_cell_margins(table)


def set_cell_text(cell, text: str, bold=False, color: RGBColor | None = None, size: float = 9.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color or INK
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Website Analytics Public How To Guide | Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Table of contents will update when opened in Microsoft Word."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(placeholder)
    run._r.append(fld_end)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = TITLE_BLUE
    title.paragraph_format.space_after = Pt(8)

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(13)
    subtitle.font.color.rgb = MUTED
    subtitle.paragraph_format.space_after = Pt(16)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, ACCENT, 18, 10),
        ("Heading 2", 13, ACCENT, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def paragraph(doc: Document, text: str = "", style: str | None = None):
    return doc.add_paragraph(text, style=style)


def bullet(doc: Document, text: str) -> None:
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.line_spacing = 1.25
    para.add_run(text)


def create_decimal_numbering(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    existing_abstract_ids = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
        if el.get(qn("w:abstractNumId")) is not None
    ]
    existing_num_ids = [
        int(el.get(qn("w:numId")))
        for el in numbering.findall(qn("w:num"))
        if el.get(qn("w:numId")) is not None
    ]
    abstract_id = (max(existing_abstract_ids) + 1) if existing_abstract_ids else 1
    num_id = (max(existing_num_ids) + 1) if existing_num_ids else 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    lvl.append(start)
    lvl.append(num_fmt)
    lvl.append(lvl_text)
    lvl.append(lvl_jc)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    lvl_override.append(start_override)
    num.append(lvl_override)
    numbering.append(num)
    return num_id


def number_with_id(doc: Document, text: str, num_id: int) -> None:
    para = doc.add_paragraph()
    p_pr = para._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.line_spacing = 1.25
    para.add_run(text)


def numbered_list(doc: Document, items: list[str]) -> None:
    for idx, item in enumerate(items, start=1):
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.25)
        para.paragraph_format.first_line_indent = Inches(-0.25)
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.line_spacing = 1.25
        run = para.add_run(f"{idx}. ")
        run.bold = True
        para.add_run(item)


def code_block(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    cell.text = ""
    for idx, line in enumerate(text.splitlines()):
        para = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(40, 47, 57)


def callout(doc: Document, title: str, body: str, fill: str = LIGHT_BLUE, color: RGBColor = DARK_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(title)
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(10.5)
    para2 = cell.add_paragraph()
    para2.paragraph_format.space_after = Pt(0)
    run2 = para2.add_run(body)
    run2.font.size = Pt(10)
    run2.font.color.rgb = INK


def add_image(doc: Document, filename: str, caption: str, width_in: float = 6.2) -> None:
    image_path = ASSETS_DIR / filename
    if not image_path.exists():
        missing = doc.add_paragraph()
        missing.add_run(f"[Missing screenshot asset: {filename}]").italic = True
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(image_path), width=Inches(width_in))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    for r in cap.runs:
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = MUTED


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(hdr[idx], header, bold=True, color=DARK_BLUE)
        set_cell_shading(hdr[idx], TABLE_HEADER)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)


def add_status_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    add_table(doc, ["Item", "Purpose", "Public-source instruction"], [[a, b, c] for a, b, c in rows], [1.55, 2.35, 2.6])


def add_appendix(doc: Document) -> None:
    doc.add_heading("Appendix A - Source files used to build this guide", level=1)
    paragraph(
        doc,
        "This guide was built from the current project source, forms, schema, README, help topics, and documentation generators. "
        "The most relevant files for future maintenance are listed below.",
    )
    for item in [
        "README.md",
        "WebsiteAnalytics.dpr and WebsiteAnalytics.dproj",
        "WebsiteAnalytics.MainForm.pas/.fmx",
        "WebsiteAnalytics.SettingsForm.pas/.fmx",
        "WebsiteAnalytics.PropertyManagerForm.pas/.fmx",
        "WebsiteAnalytics.AuthenticationDataModule.pas/.dfm",
        "WebsiteAnalytics.GA4DataModule.pas/.dfm",
        "WebsiteAnalytics.SettingsDataModule.pas/.dfm",
        "WebsiteAnalytics.AnalyticsMemoryDataModule.pas/.dfm",
        "WebsiteAnalytics.Models.pas",
        "databases/WebsiteAnalytics.schema.sql",
        "help/WebsiteAnalytics Help.hdc.json",
        "help/topics/*.html and help/assets/*.png",
        "tools/build_phase1.cmd",
        "tools/build_public_how_to_guide.py",
        "tools/refresh_websiteanalytics_docs.py",
    ]:
        bullet(doc, item)


def build_document() -> None:
    GUIDES_DIR.mkdir(parents=True, exist_ok=True)
    PDF_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    configure_styles(doc)
    add_page_number(doc.sections[0])

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Website Analytics Public How To Guide")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Build, configure, authenticate, run, and maintain the GA4 desktop dashboard from scratch")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Last changed: {LAST_CHANGED}\n")
    meta.add_run("Written for a public-source repository. Each user supplies their own Google OAuth client and GA4 property IDs.")
    add_image(doc, "real-dashboard-current-redacted.png", "Current dashboard layout with sanitized example data.", 6.4)
    callout(
        doc,
        "Public-source safety rule",
        "Never publish your local SQLite database, OAuth client secret, encrypted refresh token, credential JSON files, API keys, PEM files, or .env files. Every user should bring their own Google Cloud OAuth desktop client and their own GA4 property access.",
        LIGHT_RED,
        RED,
    )

    doc.add_page_break()
    doc.add_heading("Table of contents", level=1)
    add_toc(doc.add_paragraph())
    doc.add_page_break()

    doc.add_heading("1. What this application is", level=1)
    paragraph(
        doc,
        "Website Analytics is a Delphi FireMonkey desktop dashboard for Google Analytics 4. It was built to answer practical website questions without forcing the user through the standard GA4 web console. The program connects to GA4 through Google's Data API, retrieves selected reports, formats the returned rows in memory, and displays them in a cleaner desktop UI.",
    )
    paragraph(
        doc,
        "The application is intentionally focused. It is not a replacement for every GA4 screen. It is a daily operational dashboard: who visited, where they came from, what pages or downloads they used, what the selected date range looks like, and what is happening right now.",
    )
    add_status_table(
        doc,
        [
            ["Desktop app", "Runs as a Windows Delphi/FM X executable.", "Build from source with RAD Studio or an equivalent Delphi toolchain."],
            ["GA4 reader", "Reads data through Google Analytics Data API.", "Each user must authorize their own Google account."],
            ["Memory-only reports", "Fetched analytics rows are not stored as history.", "Close the app and the report rows are gone."],
            ["Portable settings", "SQLite stores settings beside the project/app.", "The generated local database must stay out of Git."],
        ],
    )

    doc.add_heading("2. What the current dashboard shows", level=1)
    add_image(doc, "real-dashboard-current-redacted.png", "Dashboard overview: locations, pages/downloads/actions, trend graph, realtime panel, and KPI cards.", 6.4)
    paragraph(doc, "The current UI is built around a compact dashboard, not a collection of GA4-style pages. The important parts are:")
    bullet(doc, "Hero section: app title, logo, website property selector, date range selector, auto update, connection status, Update, Settings, and Help.")
    bullet(doc, "Location grid: country, region/state, city, users, and sessions for the selected property and date range.")
    bullet(doc, "Pages/downloads/actions panel: page titles, paths, events, users, and event/download totals for the selected range.")
    bullet(doc, "Users graph: proportional bar graph for the selected range, with tooltips for exact values.")
    bullet(doc, "Realtime panel: current active users, current location if GA4 returns it, and last activity time/location.")
    bullet(doc, "KPI cards: users today, users for selected range, views today, downloads for selected range, active now, last location, top page, views last 30 minutes, events per session, and scrolled count.")

    doc.add_heading("3. What the app stores and what it does not store", level=1)
    paragraph(doc, "This distinction matters for public-source users and for privacy.")
    add_table(
        doc,
        ["Stored item", "Where", "Why", "Commit to Git?"],
        [
            ["Website definitions", "SQLite settings database", "Lets the app know which GA4 properties to query.", "No, local database only."],
            ["Dashboard defaults", "SQLite settings database", "Default site, date range, auto-refresh choice, saved grid/card preferences.", "No, local database only."],
            ["OAuth client ID/secret", "SQLite settings database", "Convenience for the local user's own Google Cloud OAuth client.", "No."],
            ["Encrypted refresh token", "SQLite settings database", "Allows silent reconnect without browser login every launch.", "No."],
            ["GA4 report rows", "Memory only", "Displayed, graphed, then discarded when the app closes.", "Never stored."],
            ["Schema", "databases/WebsiteAnalytics.schema.sql", "Creates expected settings tables and defaults.", "Yes."],
        ],
        [1.65, 1.65, 2.35, 0.85],
    )
    code_block(
        doc,
        "Tracked:\n"
        "  databases\\WebsiteAnalytics.schema.sql\n\n"
        "Ignored local files:\n"
        "  databases\\*.sqlite3\n"
        "  databases\\*.sqlite3-journal\n"
        "  databases\\*.sqlite3-wal\n"
        "  databases\\*.sqlite3-shm",
    )
    callout(
        doc,
        "Refresh token encryption",
        "The app encrypts the saved refresh token with Windows DPAPI. That is suitable for a desktop convenience token, but it does not make the local database a public artifact. Do not commit it.",
        LIGHT_GOLD,
        GOLD,
    )

    doc.add_heading("4. Prerequisites", level=1)
    add_table(
        doc,
        ["Requirement", "Details"],
        [
            ["Windows", "The current app targets Windows desktop builds."],
            ["RAD Studio / Delphi with FireMonkey", "Open and edit the .fmx forms in the IDE. Keep forms editable."],
            ["RAD Studio toolchain", r"The project scripts expect the Embarcadero toolchain under C:\Program Files (x86)\Embarcadero\Studio\37.0\bin or a compatible version."],
            ["Git", "Required for clone, branch, commit, and push workflows."],
            ["GitHub account", "Required only if you fork, push, or contribute changes."],
            ["Google account with GA4 access", "The signed-in Google account must have permission to read the GA4 properties you enter."],
            ["Google Cloud project", "Required to enable the Google Analytics Data API and create an OAuth desktop client."],
            ["GA4 properties already collecting data", "The app reads GA4; it does not install tracking tags on websites."],
        ],
        [2.1, 4.4],
    )

    doc.add_heading("5. Clone or download the source", level=1)
    paragraph(doc, "Use Git if you plan to make changes or receive updates.")
    code_block(doc, f"git clone {PUBLIC_REPO_URL}\ncd WebsiteAnalytics-GA4-Dashboard")
    paragraph(doc, "If you download the ZIP from GitHub, extract it to a normal local development folder such as:")
    code_block(doc, EXAMPLE_PROJECT_FOLDER)
    callout(
        doc,
        "Folder recommendation",
        "Avoid synced folders and restricted paths while building Delphi projects. A simple local folder makes RAD Studio, MSBuild, SQLite, and generated output easier to reason about.",
    )

    doc.add_heading("6. Repository layout", level=1)
    add_table(
        doc,
        ["Path", "Purpose"],
        [
            ["WebsiteAnalytics.dpr / .dproj", "Application entry point and RAD Studio project."],
            ["WebsiteAnalytics.MainForm.pas/.fmx", "Main editable FMX dashboard, update logic, chart, cards, grids, status, and help/settings launch points."],
            ["WebsiteAnalytics.GA4DataModule.pas/.dfm", "GA4 report request construction, API calls, and response parsing."],
            ["WebsiteAnalytics.AuthenticationDataModule.pas/.dfm", "Google OAuth desktop flow, loopback listener, token exchange, refresh-token handling."],
            ["WebsiteAnalytics.SettingsDataModule.pas/.dfm", "Portable SQLite settings creation, defaults, reads, and writes."],
            ["WebsiteAnalytics.SettingsForm.pas/.fmx", "OAuth setup, defaults, redirect port, connection/disconnect actions."],
            ["WebsiteAnalytics.PropertyManagerForm.pas/.fmx", "Website property definitions and GA4 numeric property IDs."],
            ["WebsiteAnalytics.Models.pas", "Shared data models used between data modules and UI."],
            ["databases/WebsiteAnalytics.schema.sql", "Tracked database schema for fresh settings database creation."],
            ["docs/guides", "Editable Word guides."],
            ["docs/pdf", "Companion PDF guides."],
            ["help", "Local HTML help and Help Doc Creator project file."],
            ["tools", "Build, documentation, and maintenance scripts."],
        ],
        [2.35, 4.15],
    )

    doc.add_heading("7. Build the app", level=1)
    paragraph(doc, "You can build from RAD Studio or from the command line.")
    doc.add_heading("7.1 Build in RAD Studio", level=2)
    numbered_list(
        doc,
        [
            "Open WebsiteAnalytics.dproj.",
            "Choose the desired platform and configuration, such as Win32 Debug or Win64 Debug.",
            "Build the project.",
            "Run the executable once to smoke-test form loading and startup behavior.",
        ],
    )
    doc.add_heading("7.2 Build from command line", level=2)
    paragraph(doc, "The included project build script sets up the Embarcadero environment and builds the project.")
    code_block(doc, r"tools\build_phase1.cmd")
    paragraph(doc, "Expected debug output locations:")
    code_block(doc, r"bin\Win32\Debug\WebsiteAnalytics.exe" + "\n" + r"bin\Win64\Debug\WebsiteAnalytics.exe")
    callout(
        doc,
        "Smoke-test after .fmx edits",
        "A compile can succeed while an FMX form stream still fails at runtime. If you change .fmx files, start the EXE and verify it reaches the main window without EReadError exceptions.",
        LIGHT_GOLD,
        GOLD,
    )

    doc.add_heading("8. First run and SQLite setup", level=1)
    paragraph(doc, "On first run, the app creates the portable settings database if it is missing. The repository should ship the schema, not a real user database.")
    numbered_list(
        doc,
        [
            "Start the app after a successful build.",
            "If the databases folder exists but WebsiteAnalytics.sqlite3 is missing, the app creates it.",
            "If required settings tables are missing, the app creates or updates them.",
            "Open Settings and Properties to enter your own OAuth and GA4 details.",
        ],
    )
    paragraph(doc, "A fresh public-source checkout should not already know your Google credentials or personal website properties. That is intentional.")
    doc.add_heading("8.1 What should happen on first run", level=2)
    paragraph(
        doc,
        "A clean first run should reach the main dashboard without requiring a database file to already exist. The Settings data module is responsible for creating the database and ensuring the required settings tables exist. If startup fails before the main window appears, the first thing to check is whether a streamed form property, resource, or missing runtime dependency is preventing the executable from loading.",
    )
    add_table(
        doc,
        ["First-run item", "Expected result", "If it fails"],
        [
            ["databases folder", "Created if missing.", "Confirm the application has write access to the project/app folder."],
            ["WebsiteAnalytics.sqlite3", "Created locally if missing.", "Check SQLite/FireDAC availability and folder permissions."],
            ["Default settings", "Seeded from code/schema defaults.", "Check SettingsDataModule initialization."],
            ["Generated analytics rows", "None stored.", "This is correct; report rows are fetched only during updates."],
        ],
        [1.6, 2.15, 2.75],
    )

    doc.add_heading("9. Find your numeric GA4 property IDs", level=1)
    paragraph(doc, "Every website you add needs its numeric GA4 property ID. This is one of the most common setup mistakes, so check it carefully.")
    add_table(
        doc,
        ["Thing", "Example format", "Where it belongs"],
        [
            ["GA4 property ID", "123456789", "Properties screen. This is the number the GA4 Data API uses."],
            ["Measurement ID", "G-ABCD123456", "Your website tracking tag, not this desktop app's property field."],
            ["Google Cloud project number", "101177308157", "Cloud Console only. Do not enter it as a GA4 property ID."],
            ["OAuth desktop Client ID", "101...apps.googleusercontent.com", "Settings screen OAuth Client ID field."],
            ["OAuth client secret", "Secret string from Google", "Settings screen OAuth secret field, hidden by default."],
            ["Gmail address", "name@example.com", "Only the account you sign in with; it is not an app setting."],
        ],
        [1.8, 2.1, 2.6],
    )
    numbered_list(
        doc,
        [
            "Open Google Analytics.",
            "Select the account/property for the website.",
            "Open Admin.",
            "Open Property details or Property settings.",
            "Copy the numeric Property ID.",
            "Repeat for each website you want to add.",
        ],
    )

    doc.add_heading("10. Add website properties in the app", level=1)
    add_image(doc, "real-properties-redacted.png", "Manage Properties screen with sanitized sample values.", 6.2)
    numbered_list(
        doc,
        [
            "Click Settings or Properties, depending on the current UI entry point.",
            "Open the property manager.",
            "Create or edit a property.",
            "Enter a friendly display name.",
            "Enter the public website address.",
            "Enter the numeric GA4 property ID.",
            "Enable the property if it should appear in the dashboard selector.",
            "Save the property.",
        ],
    )
    paragraph(doc, "For public screenshots and examples, use placeholder property IDs. Do not publish your real property IDs unless you explicitly choose to.")
    doc.add_heading("10.1 Adding more websites later", level=2)
    paragraph(
        doc,
        "The application was designed to start with a small set of known properties and then grow. Adding another website should not require code changes. Add a new property definition, enter the numeric GA4 property ID, enable it, and save. The dashboard selector should then include the property after the settings/property list is refreshed.",
    )
    add_table(
        doc,
        ["Field", "What to enter", "Example for public docs"],
        [
            ["Display name", "Friendly name shown in the dashboard selector.", "Example Site"],
            ["Website URL", "The public site address.", "https://example.github.io"],
            ["GA4 property ID", "The numeric GA4 property ID.", "123456789"],
            ["Enabled", "Whether this property appears in the dashboard.", "Checked"],
            ["Color", "Optional UI/report color for the property.", "Any non-private color value"],
        ],
        [1.5, 3.1, 1.9],
    )

    doc.add_heading("11. Create your Google Cloud OAuth setup", level=1)
    paragraph(doc, "A public repository should not ship a shared OAuth client. Each user should create their own Google Cloud project and OAuth desktop client.")
    doc.add_heading("11.1 Create or choose a Google Cloud project", level=2)
    numbered_list(
        doc,
        [
            "Open Google Cloud Console.",
            "Create a new project or choose an existing project for your own desktop tools.",
            "Use a plain name such as Website Analytics Desktop.",
            "Confirm you are signed in with the Google account you intend to use.",
        ],
    )
    doc.add_heading("11.2 Enable Google Analytics Data API", level=2)
    numbered_list(
        doc,
        [
            "Open APIs & Services.",
            "Open Library.",
            "Search for Google Analytics Data API.",
            "Open the API page.",
            "Click Enable.",
        ],
    )
    paragraph(doc, "If this API is not enabled, the app may authenticate successfully but fail to update with a 403 API-not-enabled error.")
    doc.add_heading("11.3 Configure OAuth consent", level=2)
    numbered_list(
        doc,
        [
            "Open APIs & Services > OAuth consent screen.",
            "Choose External unless you are intentionally using an internal Google Workspace-only app.",
            "Enter the app name, support email, and developer contact information requested by Google.",
            f"Use the read-only Analytics scope when requested: {GA4_READONLY_SCOPE}",
            "If the app is in Testing mode, add your Google account as a test user.",
            "Save the consent configuration.",
        ],
    )
    callout(
        doc,
        "Testing versus verified app",
        "For personal use, Testing mode is usually enough if your Google account is listed as a test user. For distribution to other people, do not ask them to use your private testing client. They should create their own client, or the project owner should pursue Google's OAuth app verification process.",
        LIGHT_GOLD,
        GOLD,
    )
    doc.add_heading("11.4 Create an OAuth 2.0 Desktop client", level=2)
    numbered_list(
        doc,
        [
            "Open APIs & Services > Credentials.",
            "Click Create credentials.",
            "Choose OAuth client ID.",
            "Choose Desktop app as the application type.",
            "Name it something recognizable.",
            "Create it.",
            "Copy the Client ID and Client secret.",
        ],
    )
    paragraph(doc, "The Client ID usually ends with apps.googleusercontent.com. The Client secret normally begins with a Google-generated secret prefix. They must come from the same OAuth client.")
    doc.add_heading("11.5 OAuth values people often confuse", level=2)
    paragraph(
        doc,
        "OAuth setup is the part most likely to confuse new users because Google uses several similar-looking identifiers. The desktop app needs exactly two OAuth values in Settings: the OAuth desktop Client ID and the OAuth desktop Client secret. Those are not the same as the Google account email, GA4 property ID, Measurement ID, or Google Cloud project number.",
    )
    add_table(
        doc,
        ["Question", "Correct answer"],
        [
            ["Is my Gmail address the OAuth Client ID?", "No. Your Gmail address is only the account you sign in with."],
            ["Can I use the Google Cloud project number as the GA4 property ID?", "No. The project number belongs to Google Cloud. The GA4 property ID comes from Google Analytics property settings."],
            ["Can I use the Measurement ID that starts with G-?", "No. That ID goes into website tracking code. The desktop app uses the numeric GA4 property ID."],
            ["Can I leave the OAuth secret blank?", "No. The desktop client secret must match the OAuth Client ID."],
            ["Can the public repository contain one shared secret?", "It should not. Each user should create their own OAuth desktop client."],
        ],
        [2.35, 4.15],
    )

    doc.add_heading("12. Configure Google sign-in in Website Analytics", level=1)
    add_image(doc, "real-settings-redacted.png", "Settings and Google sign-in screen with private values redacted.", 6.2)
    paragraph(doc, "The default redirect URI is:")
    code_block(doc, DEFAULT_REDIRECT_URI)
    paragraph(doc, "The app uses a loopback listener on the local computer. During sign-in, Google redirects the browser back to this local URI so the desktop app can receive the authorization code.")
    numbered_list(
        doc,
        [
            "Open Settings.",
            "Enter your OAuth desktop Client ID.",
            "Use Show only when you need to enter or edit the OAuth Client secret.",
            "Enter the OAuth Client secret from the same desktop client.",
            "Leave the loopback port at 53682 unless another local program uses that port.",
            "Confirm the Redirect URI shown by the app.",
            "Click Connect Google.",
            "Complete the browser sign-in with the Google account that can read the GA4 properties.",
            "If Google shows a testing/unverified warning and this is your own OAuth client, continue intentionally.",
            "Return to the app after the browser says you may close the tab.",
            "Click Save and Close.",
        ],
    )
    callout(
        doc,
        "What is saved after sign-in",
        "The app stores settings locally and saves the refresh token encrypted with Windows DPAPI so later startups can reconnect silently. It does not save GA4 report rows.",
    )
    doc.add_heading("12.1 Settings fields explained", level=2)
    add_table(
        doc,
        ["Setting", "Meaning", "Public-source guidance"],
        [
            ["Default website", "The property selected when the dashboard opens.", "Use a placeholder/default in public docs; users choose their own."],
            ["Default date range", "Initial reporting window.", "Last 7 days is a sensible first setting."],
            ["Auto update", "Automatically refreshes on the timer while connected.", "Useful when the app is left running all day."],
            ["OAuth Client ID", "Desktop OAuth client identifier.", "Each user supplies their own."],
            ["OAuth Client secret", "Secret paired with the Client ID.", "Masked in the UI; do not publish screenshots showing it."],
            ["Loopback port", "Local HTTP port used for browser redirect.", "Default 53682 unless there is a local conflict."],
            ["Redirect URI", "Local URL Google redirects back to.", f"Default is {DEFAULT_REDIRECT_URI}."],
        ],
        [1.5, 2.45, 2.55],
    )

    doc.add_heading("13. Silent login and token refresh", level=1)
    paragraph(doc, "The first authorization needs the browser. After that, the app should normally avoid the repeated Google login/continue screens.")
    add_table(
        doc,
        ["Token or action", "What happens"],
        [
            ["Access token", "Short-lived token kept in memory and used for GA4 API calls."],
            ["Refresh token", "Longer-lived token encrypted with Windows DPAPI and stored in SQLite."],
            ["Startup reconnect", "The app decrypts the refresh token and asks Google for a fresh access token silently."],
            ["Auto update", "The 60-second timer can refresh dashboard data without browser prompts while authorization remains valid."],
            ["Expired or revoked token", "The app falls back to visible Google sign-in."],
            ["Disconnect", "Clears the in-memory token and saved encrypted refresh token."],
        ],
        [1.7, 4.8],
    )

    doc.add_heading("14. Run your first dashboard update", level=1)
    add_image(doc, "real-hero-redacted.png", "Hero section: property selector, date range, auto update, status, Update, Settings, and Help.", 6.2)
    numbered_list(
        doc,
        [
            "Choose one website property from the Website property selector.",
            "Choose the date range you want to inspect.",
            "Click Update.",
            "Read the bottom status line if the update fails or returns no data.",
            "Hover over graph bars for exact values.",
            "Leave Auto 60 sec checked if you want the app to keep refreshing while it remains open.",
        ],
    )
    paragraph(doc, "When All websites is selected, the app queries enabled properties and combines rows in memory. Individual websites are usually clearer when you are diagnosing a specific site's behavior.")
    doc.add_heading("14.1 What a successful update does", level=2)
    paragraph(
        doc,
        "A successful update is more than a single GA4 call. The app validates authentication, refreshes the access token if needed, resolves the selected property or the enabled property list, runs the requested GA4 reports, merges or normalizes rows where needed, updates the in-memory snapshot, and refreshes the visible grids, cards, graph, realtime panel, and status bar.",
    )
    add_table(
        doc,
        ["Step", "What happens"],
        [
            ["Authentication check", "Uses the current access token or silently refreshes it with the encrypted refresh token."],
            ["Property resolution", "Uses the selected property or builds the enabled-property list for All websites."],
            ["Historical reports", "Requests summary, trend, location, and content/action rows for the selected date range."],
            ["Realtime reports", "Requests active-now and recent location/activity information from GA4 realtime reporting."],
            ["Normalization", "Fills missing dates with zeroes so graphs remain proportional and readable."],
            ["Presentation", "Updates the grids, cards, graph, realtime text, and status bar."],
        ],
        [1.7, 4.8],
    )

    doc.add_heading("15. Understand the reports", level=1)
    add_table(
        doc,
        ["Screen area", "Question answered", "Typical GA4 source"],
        [
            ["Location grid", "Where did visitors come from?", "country, region, city with users and sessions."],
            ["Pages/downloads/actions", "What content or actions did visitors use?", "page title, page path, event name, users, views, event counts, downloads."],
            ["Users graph", "How did users trend over the selected period?", "date plus users, normalized so empty dates show as zero."],
            ["Realtime panel", "What is happening right now?", "runRealtimeReport active users and current location when GA4 returns it."],
            ["KPI cards", "What are the headline values?", "summary, realtime, content, and engagement metrics."],
        ],
        [1.75, 2.2, 2.55],
    )
    add_image(doc, "real-locations-grid-redacted.png", "Location grid: country, region/state, city, users, and sessions.", 6.2)
    add_image(doc, "real-pages-downloads-redacted.png", "Pages/downloads/actions panel: page title, path, action/event, users, and events/downloads.", 6.2)
    add_image(doc, "real-users-graph.png", "Users graph: selected range, whole-number scale, and hover tooltip values.", 6.2)
    add_image(doc, "real-realtime-panel.png", "Realtime panel: current active visitors, current location, and last activity when available.", 6.2)
    add_image(doc, "real-summary-cards.png", "KPI card area with headline values.", 6.2)
    doc.add_heading("15.1 KPI card meanings", level=2)
    add_table(
        doc,
        ["Card", "Meaning", "Notes"],
        [
            ["Users today", "Users returned by GA4 for today's date.", "Standard report data may lag behind realtime."],
            ["Users for selected range", "Users for the chosen date range.", "Changes when date range changes."],
            ["Views today", "Screen/page views for today.", "Useful for quick activity level."],
            ["Downloads", "Download-related events for the selected range.", "Depends on GA4 receiving download events."],
            ["Active now", "Realtime active users.", "Comes from GA4 realtime endpoint."],
            ["Last Location", "Most recent usable location returned by realtime/location reports.", "GA4 may return None or (not set)."],
            ["Top page", "Highest page/action item from current reports.", "Home page is shown when path is /."],
            ["Views last 30 min", "Realtime or near-realtime views in the last 30 minutes.", "Can differ from standard report views."],
            ["Events/session", "Event count divided by sessions where data is available.", "Shows activity density."],
            ["Scrolled", "Scroll events counted for selected range.", "Depends on GA4 event collection."],
        ],
        [1.45, 2.55, 2.5],
    )
    doc.add_heading("15.2 Handling no data and '(not set)'", level=2)
    paragraph(
        doc,
        "GA4 can legitimately return empty rows, zeroes, or '(not set)' values. The app should display those conditions plainly instead of inventing data. A blank result for one date does not necessarily mean the app is broken; it may mean GA4 has no rows for that date/property/report combination, the selected range is too narrow, or Google's processing has not caught up yet.",
    )
    bullet(doc, "Use a wider date range if you expect activity but see no rows.")
    bullet(doc, "Test one property at a time before judging All websites.")
    bullet(doc, "Compare realtime and standard reports with care; they are different GA4 endpoints.")
    bullet(doc, "Treat '(not set)' as a GA4-returned value, not as a local application label.")

    doc.add_heading("16. GA4 data requested by the app", level=1)
    paragraph(doc, "The current data module asks GA4 for several focused reports, then translates them into the dashboard snapshot.")
    add_table(
        doc,
        ["Purpose", "Dimensions", "Metrics", "Used for"],
        [
            ["Summary", "None", "activeUsers, sessions, screenPageViews, engagementRate", "Headline cards."],
            ["Trend", "date", "activeUsers, sessions, screenPageViews, engagementRate, eventCount", "Bar graph over selected range."],
            ["Geography", "country, region, city", "activeUsers, sessions, engagementRate", "Location grid and last-location style values."],
            ["Content/actions", "pagePath, pageTitle, eventName", "screenPageViews, activeUsers, eventCount, userEngagementDuration", "Top pages/downloads/actions panel and cards."],
            ["Realtime", "country, region, city where available", "activeUsers, screenPageViews where available", "Active now, current location, views last 30 minutes."],
            ["Support reports", "source/medium/campaign, device, language", "users, sessions, engagement-related metrics", "Available for future dashboard expansion."],
        ],
        [1.35, 1.55, 2.35, 1.25],
    )
    callout(
        doc,
        "Realtime can differ from standard reports",
        "GA4 realtime data is not always synchronized with normal report data. The app tries to keep the dashboard internally consistent, but Google's realtime endpoint and standard reporting endpoint do not always update on the same schedule.",
        LIGHT_GOLD,
        GOLD,
    )

    doc.add_heading("17. Troubleshooting", level=1)
    add_table(
        doc,
        ["Symptom", "Likely cause", "What to do"],
        [
            ["Google login appears every launch", "No usable saved refresh token.", "Reconnect once from Settings. If it repeats, check whether access was revoked or the database moved to another Windows user/machine."],
            ["Access blocked / unverified app", "OAuth app is in Testing or not verified.", "For private use, add yourself as a test user. For public distribution, create your own OAuth client or pursue verification."],
            ["Token exchange failed", "Client secret is missing, wrong, or mismatched.", "Copy the secret from the same OAuth desktop client as the Client ID."],
            ["403 API not enabled", "Google Analytics Data API is disabled in the Cloud project.", "Enable Google Analytics Data API in APIs & Services > Library."],
            ["403 insufficient permissions", "Wrong property ID or signed-in Google account cannot read the property.", "Verify the numeric GA4 property ID and GA4 Admin/User access."],
            ["No rows for a date", "No data returned for that property/range.", "Try a wider date range or check GA4 web console."],
            ["Downloads missing", "GA4 is not emitting file_download or custom download events.", "Check Enhanced Measurement, outbound/download tracking, and date range."],
            ["All websites looks low", "One property may fail, no-data may be returned, or rows may be grouped differently.", "Test each property individually, then retry All websites and read the status line."],
            ["FMX EReadError", "Unsupported property was streamed into .fmx.", "Remove the invalid property and smoke-test startup."],
            ["Shortcut says EXE has no icons", "Icon was not embedded into that build target.", "Check RAD Studio project icon/resource settings for Win32/Win64 and rebuild the exact EXE you inspect."],
        ],
        [1.55, 2.15, 2.8],
    )
    doc.add_heading("17.1 A practical troubleshooting order", level=2)
    numbered_list(
        doc,
        [
            "Check the status bar message first. It usually tells whether the problem is authentication, permission, API access, parsing, or no data.",
            "Open Settings and confirm the OAuth Client ID and Client secret belong to the same Desktop app credential.",
            "Open Google Cloud Console and confirm Google Analytics Data API is enabled.",
            "Open Google Analytics Admin and confirm the signed-in Google account can read the numeric property ID.",
            "Try one known-active property instead of All websites.",
            "Try Last 90 days instead of Today if you are not sure recent data exists.",
            "Open Diagnostics if the app has a request/parser message.",
            "If the app will not start, look for .fmx/.dfm EReadError messages and fix unsupported streamed properties.",
        ],
    )
    doc.add_heading("17.2 Error wording to recognize", level=2)
    add_table(
        doc,
        ["Message pattern", "Meaning"],
        [
            ["client_secret is missing", "The token exchange reached Google, but the local OAuth secret value is missing or not being sent."],
            ["Access blocked / not verified", "Google OAuth consent is not verified or the user is not an allowed test user."],
            ["Data API has not been used", "The Google Analytics Data API is not enabled in the selected Google Cloud project."],
            ["User does not have sufficient permissions", "The signed-in Google account cannot read the requested GA4 property, or the property ID is wrong."],
            ["Value 'rows' not found", "The parser expected rows but GA4 returned an empty or different-shaped response. The app should handle this safely."],
            ["EReadError reading form", "The executable cannot load a streamed Delphi form/resource property."],
        ],
        [2.35, 4.15],
    )

    doc.add_heading("18. GitHub setup for contributors", level=1)
    paragraph(doc, "A user who only downloads the app does not need GitHub write access. A contributor or maintainer needs Git identity and authentication.")
    doc.add_heading("18.1 Configure Git identity", level=2)
    code_block(doc, 'git config --global user.name "Your Name"\ngit config --global user.email "you@example.com"')
    doc.add_heading("18.2 Authenticate to GitHub", level=2)
    paragraph(doc, "Use Git Credential Manager, GitHub CLI, or SSH. GitHub CLI is explicit and easy to verify:")
    code_block(doc, "gh auth login\n# Choose GitHub.com, HTTPS or SSH, then complete the browser/device login.")
    doc.add_heading("18.3 Normal contribution workflow", level=2)
    code_block(
        doc,
        f"git clone {PUBLIC_REPO_URL}\n"
        "cd WebsiteAnalytics-GA4-Dashboard\n"
        "git checkout -b my-change\n\n"
        "git status\n"
        "git add <only-the-files-you-intended>\n"
        "git commit -m \"Describe the change\"\n"
        "git push origin my-change",
    )
    callout(
        doc,
        "Before every push",
        "Run git status. Confirm you are not staging databases/*.sqlite3, OAuth client secrets, encrypted refresh tokens, credential JSON files, .env files, keys, PEM files, or private screenshots.",
        LIGHT_RED,
        RED,
    )

    doc.add_heading("19. Public release safety checklist", level=1)
    for item in [
        "Keep databases/*.sqlite3 and SQLite sidecar files ignored.",
        "Keep databases/WebsiteAnalytics.schema.sql tracked.",
        "Do not commit OAuth client secrets, refresh tokens, generated databases, credential JSON files, keys, PEM files, or .env files.",
        "Use sanitized screenshots in README, guides, help, issues, and release notes.",
        "Do not publish private property IDs unless you intentionally decide they are safe to disclose.",
        "Explain that the OAuth Client ID is not a Gmail address.",
        "Explain that the GA4 Property ID is not the Measurement ID and not the Google Cloud project number.",
        "If a secret was ever committed, revoke it in Google Cloud before making the repository public.",
        "If a secret was committed in history, consider cleaning Git history before public release.",
        "Expect GitHub secret scanning or safety review to flag OAuth-looking credentials in source, docs, screenshots, or commit history.",
    ]:
        bullet(doc, item)

    doc.add_heading("20. Google verification for broader public use", level=1)
    paragraph(
        doc,
        "For private use, each user creating their own OAuth desktop client is usually the simplest path. For a widely distributed application with a shared OAuth client, Google may require OAuth app verification.",
    )
    numbered_list(
        doc,
        [
            "Use the least-privileged scope required, normally analytics.readonly.",
            "Prepare a clear app name, support email, privacy policy, and user-facing explanation of data use.",
            "Explain that the app reads GA4 report data and does not store report rows.",
            "Provide screenshots or a demonstration video if Google requests them.",
            "Submit the OAuth consent app for verification from Google Cloud Console.",
            "Do not distribute a testing-mode OAuth client to general public users.",
        ],
    )
    doc.add_heading("20.1 Practical recommendation for this project", level=2)
    paragraph(
        doc,
        "For an open-source desktop utility, the lowest-friction public model is to publish the source code and require each user to create their own Google Cloud OAuth desktop client. That avoids shipping a shared secret and avoids asking strangers to trust a testing-mode OAuth app owned by somebody else. If the project later becomes a packaged product with non-technical users, then a verified OAuth app and a polished privacy policy become much more important.",
    )
    doc.add_heading("20.2 What a privacy explanation should say", level=2)
    bullet(doc, "The app requests read-only GA4 access.")
    bullet(doc, "The app uses GA4 report data to display local dashboard panels.")
    bullet(doc, "The app does not upload analytics data to a third-party service.")
    bullet(doc, "The app does not store fetched analytics report rows as history.")
    bullet(doc, "The app stores local settings and an encrypted refresh token so it can reconnect.")
    bullet(doc, "Disconnect clears the saved token.")

    doc.add_heading("21. Maintenance rules for developers", level=1)
    bullet(doc, "Keep FMX forms editable in RAD Studio. Prefer Object Inspector settings for layout and visual values unless runtime behavior is truly necessary.")
    bullet(doc, "Do not add helper code unless it has a clear purpose and the project owner approves it.")
    bullet(doc, "Back up before project changes when working under the project directives.")
    bullet(doc, "After approved changes are complete and validated, commit and push the intended files.")
    bullet(doc, "Build and smoke-test after form changes, resource/icon changes, OAuth changes, GA4 parser changes, and documentation generator changes.")
    bullet(doc, "Regenerate the user guide, engineering guide, help, and this How To guide when the UI or setup flow changes.")

    doc.add_heading("22. From-zero checklist", level=1)
    for item in [
        "Clone or download the public repository.",
        "Open or build WebsiteAnalytics.dproj.",
        "Run the app once so the local settings database is created.",
        "Find the numeric GA4 property IDs for your websites.",
        "Add your website properties in the app.",
        "Create or choose a Google Cloud project.",
        "Enable Google Analytics Data API.",
        "Configure OAuth consent and add yourself as a test user if needed.",
        "Create an OAuth 2.0 Desktop client.",
        "Enter the OAuth Client ID and Client secret in Settings.",
        f"Confirm the redirect URI is {DEFAULT_REDIRECT_URI} unless you changed the port.",
        "Connect Google once through the browser.",
        "Save Settings.",
        "Choose a property and date range.",
        "Click Update.",
        "Leave Auto 60 sec enabled if you want a live dashboard.",
        "Keep your generated database and credentials out of Git.",
    ]:
        bullet(doc, item)

    add_appendix(doc)
    doc.save(DOCX_PATH)


def export_pdf_with_word() -> None:
    ps = f"""
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$word.DisplayAlerts = 0
$doc = $word.Documents.Open('{DOCX_PATH}')
$doc.TablesOfContents | ForEach-Object {{ $_.Update() }}
$doc.Fields.Update() | Out-Null
$doc.Save()
$doc.ExportAsFixedFormat('{PDF_PATH}', 17)
$doc.Close($false)
$word.Quit()
"""
    subprocess.run(
        [
            "powershell.exe",
            "-NoProfile",
            "-ExecutionPolicy",
            "Bypass",
            "-Command",
            ps,
        ],
        check=True,
    )


if __name__ == "__main__":
    build_document()
    export_pdf_with_word()
    print(DOCX_PATH)
    print(PDF_PATH)
